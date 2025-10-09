import os
import shutil
import uuid
import xml.etree.ElementTree as ET
import zipfile
from datetime import datetime
from django.conf import settings
from django.db.models import Count, Q
from django.http import FileResponse
from django.utils import timezone
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
from rest_framework import status
from rest_framework.parsers import MultiPartParser
from rest_framework.response import Response
from rest_framework.views import APIView
from docx import Document as DocxDocument
from .models import Document, Paragraph, Comment, DocumentImage, ParagraphImage
from .serializers import DocumentSerializer


class XMLFormattingMixin:
    """Mixin class providing XML formatting methods for DOCX processing"""
    
    def _write_xml_with_proper_formatting(self, tree, file_path):
        """Write XML with proper formatting to avoid corruption"""
        try:
            import xml.dom.minidom
            
            # First write to get XML content
            rough_string = ET.tostring(tree.getroot(), encoding='unicode')
            
            # Add proper XML declaration for Word compatibility
            xml_content = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n' + rough_string
            
            # Parse and prettify with minidom
            dom = xml.dom.minidom.parseString(xml_content.encode('utf-8'))
            pretty_xml = dom.toprettyxml(indent="  ", encoding='utf-8').decode('utf-8')
            
            # Clean up extra blank lines that minidom adds
            lines = [line for line in pretty_xml.split('\n') if line.strip()]
            clean_pretty_xml = '\n'.join(lines)
            
            # Ensure XML declaration is Word-compatible
            if clean_pretty_xml.startswith('<?xml version="1.0" encoding="utf-8"?>'):
                clean_pretty_xml = clean_pretty_xml.replace(
                    '<?xml version="1.0" encoding="utf-8"?>',
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                )
            elif clean_pretty_xml.startswith('<?xml version="1.0"?>'):
                clean_pretty_xml = clean_pretty_xml.replace(
                    '<?xml version="1.0"?>',
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                )
            
            # Write the properly formatted XML
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(clean_pretty_xml)
                
            # Verify the written XML is valid
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                ET.fromstring(content.encode('utf-8'))
                print(f"Successfully wrote Word-compatible XML to {os.path.basename(file_path)}")
            except ET.ParseError as e:
                print(f"Warning: Generated XML may be malformed in {file_path}: {e}")
                    
        except Exception as e:
            print(f"Error formatting XML for {file_path}: {e}")
            print("Falling back to basic XML writing...")
            # Fallback to basic writing
            tree.write(file_path, encoding='utf-8', xml_declaration=True)

    def _format_xml_file(self, file_path):
        """Format an existing XML file to have proper indentation"""
        try:
            import xml.dom.minidom
            
            # Read the existing XML
            with open(file_path, 'r', encoding='utf-8') as f:
                xml_content = f.read()
            
            # Parse and prettify
            dom = xml.dom.minidom.parseString(xml_content.encode('utf-8'))
            pretty_xml = dom.toprettyxml(indent="  ", encoding='utf-8').decode('utf-8')
            
            # Clean up extra blank lines
            lines = [line for line in pretty_xml.split('\n') if line.strip()]
            clean_pretty_xml = '\n'.join(lines)
            
            # Ensure XML declaration is Word-compatible
            if clean_pretty_xml.startswith('<?xml version="1.0" encoding="utf-8"?>'):
                clean_pretty_xml = clean_pretty_xml.replace(
                    '<?xml version="1.0" encoding="utf-8"?>',
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                )
            elif clean_pretty_xml.startswith('<?xml version="1.0"?>'):
                clean_pretty_xml = clean_pretty_xml.replace(
                    '<?xml version="1.0"?>',
                    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                )
            
            # Write back the formatted XML
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(clean_pretty_xml)
                
            return True
            
        except Exception as e:
            print(f"Error formatting {file_path}: {e}")
            return False

    def _recreate_docx_with_proper_xml_formatting(self, file_path, temp_dir):
        """Recreate DOCX ensuring all XML files are properly formatted"""
        try:
            import xml.dom.minidom
            
            print("Formatting all XML files for Word compatibility...")
            
            # Find all XML files in the temp directory
            xml_files = []
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    if file.endswith('.xml'):
                        xml_files.append(os.path.join(root, file))
            
            print(f"Found {len(xml_files)} XML files to format")
            
            # Format each XML file
            for xml_file in xml_files:
                rel_path = os.path.relpath(xml_file, temp_dir).replace('\\', '/')
                
                # Check if file was already formatted by checking if it has proper indentation
                try:
                    with open(xml_file, 'r', encoding='utf-8') as f:
                        content = f.read()
                    
                    # Check if already properly formatted (has multiple lines and indentation)
                    if content.count('\n') > 5 and ('  <' in content or '    <' in content):
                        print(f"  {rel_path}: Already properly formatted")
                        continue
                    
                    # Format the file
                    success = self._format_xml_file(xml_file)
                    if success:
                        print(f"  {rel_path}: Formatted successfully")
                    else:
                        print(f"  {rel_path}: Formatting failed, kept original")
                        
                except Exception as e:
                    print(f"  {rel_path}: Error checking/formatting - {e}")
            
            print("Recreating DOCX with formatted XML files...")
            
            # Recreate the DOCX file
            with zipfile.ZipFile(file_path, 'w', zipfile.ZIP_DEFLATED) as docx_out:
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path_full = os.path.join(root, file)
                        # Use forward slashes for archive names (ZIP standard)
                        archive_name = os.path.relpath(file_path_full, temp_dir).replace('\\', '/')
                        docx_out.write(file_path_full, archive_name)
            
            print("DOCX recreation completed with properly formatted XML")
            return True
            
        except Exception as e:
            print(f"Error recreating DOCX with proper XML formatting: {e}")
            return False

    def delete_comment_from_docx(self, file_path, comment_id):
        """Delete a comment from the DOCX file"""
        backup_path = None
        temp_dir = None
        
        try:
            # Create a backup
            backup_path = file_path + '.backup'
            shutil.copy2(file_path, backup_path)
            
            # Extract the DOCX
            temp_dir = file_path + '_temp'
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Update comments.xml
            comments_path = os.path.join(temp_dir, 'word', 'comments.xml')
            
            if os.path.exists(comments_path):
                ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
                
                tree = ET.parse(comments_path)
                root = tree.getroot()
                
                namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # Find and remove the comment
                comments = root.findall('.//w:comment', namespaces)
                for comment in comments:
                    if comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id') == str(comment_id):
                        root.remove(comment)
                        break
                
                self._write_xml_with_proper_formatting(tree, comments_path)
            
            # Remove comment references from document.xml
            self.remove_comment_references_from_document(temp_dir, comment_id)
            
            # Recreate the DOCX file with proper XML formatting for ALL files
            self._recreate_docx_with_proper_xml_formatting(file_path, temp_dir)
            
            # Cleanup
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            if backup_path and os.path.exists(backup_path):
                os.remove(backup_path)
                
        except Exception as e:
            # Restore backup if something went wrong
            if backup_path and os.path.exists(backup_path):
                shutil.copy2(backup_path, file_path)
                os.remove(backup_path)
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            raise e

    def remove_comment_references_from_document(self, temp_dir, comment_id):
        """Remove comment references from document.xml"""
        document_path = os.path.join(temp_dir, 'word', 'document.xml')
        
        if not os.path.exists(document_path):
            return
        
        ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
        
        tree = ET.parse(document_path)
        root = tree.getroot()
        
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Create parent map for efficient parent lookup
        parent_map = {child: parent for parent in root.iter() for child in parent}
        
        # Helper function to remove elements using parent map
        def remove_elements_by_xpath(xpath_pattern):
            elements_to_remove = root.findall(xpath_pattern, namespaces)
            for element in elements_to_remove:
                parent = parent_map.get(element)
                if parent is not None:
                    parent.remove(element)
        
        # Remove comment range starts
        remove_elements_by_xpath(f'.//w:commentRangeStart[@w:id="{comment_id}"]')
        
        # Remove comment range ends  
        remove_elements_by_xpath(f'.//w:commentRangeEnd[@w:id="{comment_id}"]')
        
        # Remove comment references
        remove_elements_by_xpath(f'.//w:commentReference[@w:id="{comment_id}"]')
        
        self._write_xml_with_proper_formatting(tree, document_path)


class UploadDocumentView(APIView):
    parser_classes = [MultiPartParser]

    def extract_comments_from_docx(self, file_path):
        comments_data = []
        
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                if 'word/comments.xml' not in docx_zip.namelist():
                    print("No comments found in document")
                    return comments_data
                
                comments_xml = docx_zip.read('word/comments.xml')
                root = ET.fromstring(comments_xml)
                
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                for comment in root.findall('.//w:comment', namespaces):
                    comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                    date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                    
                    comment_text = ""
                    for p in comment.findall('.//w:p', namespaces):
                        para_text = ""
                        for t in p.findall('.//w:t', namespaces):
                            if t.text:
                                para_text += t.text
                        if para_text.strip():
                            comment_text += para_text + "\n"
                    
                    comment_text = comment_text.strip()
                    
                    if comment_text:
                        paragraph_id = self.find_comment_paragraph_id(docx_zip, comment_id, namespaces)
                        
                        comments_data.append({
                            'comment_id': comment_id,
                            'author': author,
                            'text': comment_text,
                            'date': date,
                            'paragraph_id': paragraph_id
                        })
        
        except Exception as e:
            print(f"Error extracting comments: {e}")
            
        return comments_data

    def find_comment_paragraph_id(self, docx_zip, comment_id, namespaces):
        try:
            document_xml = docx_zip.read('word/document.xml')
            root = ET.fromstring(document_xml)
            
            paragraph_counter = 0
            
            for para in root.findall('.//w:p', namespaces):
                para_text = ""
                for t in para.findall('.//w:t', namespaces):
                    if t.text:
                        para_text += t.text
                
                if para_text.strip():
                    paragraph_counter += 1
                    
                    comment_refs = para.findall(f'.//w:commentReference[@w:id="{comment_id}"]', namespaces)
                    comment_starts = para.findall(f'.//w:commentRangeStart[@w:id="{comment_id}"]', namespaces)
                    
                    if comment_refs or comment_starts:
                        return paragraph_counter
            
        except Exception as e:
            print(f"Error finding comment paragraph: {e}")
        
        return 1

    def post(self, request):
        if 'file' not in request.FILES:
            return Response({
                'status': 'error',
                'message': 'No file provided',
                'code': 'no_file'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        if not file.name.endswith('.docx'):
            return Response({
                'status': 'error',
                'message': 'Only .docx files are allowed',
                'code': 'invalid_file_type'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Save file
        filename = f"{uuid.uuid4()}_{file.name}"
        file_path = os.path.join(settings.MEDIA_ROOT, filename)
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
        
        with open(file_path, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)

        try:
            # Create document instance with version 1
            document = Document.objects.create(
                filename=file.name,
                file_path=file_path,
                is_editable=True,  # Make all documents editable
                version_number=1,
                version_status='original',
                base_document=None,  # This is the original document
                parent_document=None,  # No parent for original upload
                created_from_comments=False
            )
            
            # Use enhanced parser
            from .docx_parser import EnhancedDocxParser
            parser = EnhancedDocxParser(file_path, document)
            paragraphs_data = parser.parse_document()
            
            # Build paragraph objects dictionary for comment linking
            paragraph_objects = {}
            for para_data in paragraphs_data:
                paragraph = Paragraph.objects.get(
                    document=document,
                    paragraph_id=para_data['id']
                )
                paragraph_objects[para_data['id']] = paragraph

            comments_data = []
            extracted_comments = self.extract_comments_from_docx(file_path)
            
            for comment_data in extracted_comments:
                try:
                    paragraph_id = int(comment_data['paragraph_id'])
                    paragraph = paragraph_objects.get(paragraph_id)
                    
                    if paragraph:
                        comment_obj = Comment.objects.create(
                            document=document,
                            paragraph=paragraph,
                            comment_id=int(comment_data['comment_id']),
                            author=comment_data['author'],
                            text=comment_data['text']
                        )
                        comments_data.append({
                            'id': comment_obj.comment_id,
                            'author': comment_obj.author,
                            'text': comment_obj.text,
                            'paragraph_id': paragraph_id
                        })
                        
                except Exception as e:
                    print(f"Error creating comment: {e}")
                    continue
        
            return Response({
                'status': 'success',
                'message': 'Document uploaded successfully',
                'data': {
                    'document_id': document.id,
                    'paragraphs': paragraphs_data,
                    'comments': comments_data
                }
            })
            
        except Exception as e:
            print(f"Error parsing document: {e}")
            if os.path.exists(file_path):
                os.remove(file_path)
            return Response({'error': f'Error parsing document: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class EditParagraphView(XMLFormattingMixin, APIView):
    def put(self, request):
        document_id = request.data.get('document_id')
        paragraph_id = request.data.get('paragraph_id')
        new_text = request.data.get('text', '')
        
        if not all([document_id, paragraph_id]):
            return Response({'error': 'Missing required fields'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            print(f"DEBUG: EditParagraphView.put called with document_id={document_id}, paragraph_id={paragraph_id}")
            
            # Use get_document method if available (for full editor), otherwise use direct lookup
            if hasattr(self, 'get_document'):
                document = self.get_document(document_id)
                if not document:
                    return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
            else:
                document = Document.objects.get(id=document_id)
            
            # TRACK EDITED PARAGRAPHS: Mark this paragraph as edited if it has comments
            paragraph_has_comments = Comment.objects.filter(
                document=document, 
                paragraph__paragraph_id=paragraph_id
            ).exists()
            
            version_created = False
            new_version_id = None
            new_version_number = None
            version_message = None
            
            if paragraph_has_comments and document.version_status == 'commented':
                # Mark this paragraph as edited
                document.mark_paragraph_edited(paragraph_id)
                print(f"DEBUG: Marked paragraph {paragraph_id} as edited")
                
                # Check if all commented paragraphs have been edited
                if document.all_commented_paragraphs_edited():
                    print(f"DEBUG: All commented paragraphs have been edited, creating new version...")
                    
                    import os
                    import shutil
                    
                    # Create new version automatically
                    base_doc = document.base_document or document
                    next_version_number = base_doc.get_next_version_number()
                    
                    # Create new filename
                    original_name = base_doc.filename
                    name_parts = os.path.splitext(original_name)
                    new_filename = f"{name_parts[0]}_v{next_version_number}{name_parts[1]}"
                    
                    # Copy file
                    original_path = document.file_path
                    media_dir = os.path.dirname(original_path)
                    new_file_path = os.path.join(media_dir, new_filename)
                    shutil.copy2(original_path, new_file_path)
                    
                    # Create new document version
                    new_version = Document.objects.create(
                        filename=new_filename,
                        file_path=new_file_path,
                        is_editable=True,
                        version_number=next_version_number,
                        version_status='edited',
                        base_document=base_doc,
                        parent_document=document,
                        created_from_comments=True,
                        version_notes=f'Auto-created v{next_version_number} - all commented paragraphs edited'
                    )
                    
                    # Copy paragraphs from current version
                    current_paragraphs = document.paragraphs.all().order_by('paragraph_id')
                    paragraph_mapping = {}
                    
                    for old_paragraph in current_paragraphs:
                        new_paragraph = Paragraph.objects.create(
                            document=new_version,
                            paragraph_id=old_paragraph.paragraph_id,
                            text=old_paragraph.text,
                            html_content=old_paragraph.html_content,
                            has_images=old_paragraph.has_images
                        )
                        paragraph_mapping[old_paragraph.paragraph_id] = new_paragraph
                        
                        # Copy paragraph images if any
                        for para_image in old_paragraph.paragraph_images.all():
                            ParagraphImage.objects.create(
                                paragraph=new_paragraph,
                                document_image=para_image.document_image,
                                position_in_paragraph=para_image.position_in_paragraph
                            )
                    
                    # Update original document status to archived
                    document.version_status = 'archived'
                    document.save()
                    
                    # Switch to editing the new version
                    document = new_version
                    document_id = new_version.id
                    version_created = True
                    new_version_id = new_version.id
                    new_version_number = next_version_number
                    version_message = f'All commented paragraphs edited - created v{next_version_number}'
                    
                    print(f"DEBUG: Auto-created v{next_version_number} (ID: {new_version.id}) - all commented paragraphs completed")
                else:
                    remaining = document.get_remaining_commented_paragraphs()
                    print(f"DEBUG: Still {len(remaining)} commented paragraphs to edit: {remaining}")
                    version_message = f'Progress: {len(document.edited_commented_paragraphs)}/{len(document.get_commented_paragraph_ids())} commented paragraphs edited'
            
            paragraph = Paragraph.objects.get(document=document, paragraph_id=paragraph_id)
            print(f"DEBUG: Found paragraph {paragraph_id}, processing ML compliance checks...")
            
            # SMART COMMENT MANAGEMENT: Check ML compliance before deciding to delete comments
            # New workflow: Comment → Edit → ML Check → Delete only if compliant
            comments_to_check = Comment.objects.filter(paragraph=paragraph)
            ml_results = []
            compliant_comment_ids = []
            deleted_comment_ids = []
            
            if comments_to_check.exists():
                print(f"ML compliance checking {comments_to_check.count()} comments for paragraph {paragraph_id}")
                
                # Get original text for ML comparison
                original_text = paragraph.text or ""
                
                # Check each comment for ML compliance
                for comment in comments_to_check:
                    try:
                        print(f"DEBUG: Checking compliance for comment {comment.comment_id}")
                        # Use ML or fallback to basic compliance checking
                        ml_result = self.check_comment_compliance(original_text, comment.text, new_text)
                        print(f"DEBUG: ML result for comment {comment.comment_id}: {ml_result}")
                        
                        # Determine proper status based on score (override ML prediction if needed)
                        final_status = ml_result['prediction']
                        score = ml_result['compliance_score']
                        
                        # Ensure consistent classification: only "compliant" if score >= 0.6
                        if score >= 0.6:
                            final_status = 'compliant'
                        elif score >= 0.3:
                            final_status = 'partial'
                        else:
                            final_status = 'non_compliant'
                            
                        # Update comment status based on corrected classification
                        comment.compliance_status = final_status
                        comment.compliance_score = score
                        comment.last_checked = timezone.now()
                        comment.save()
                        
                        ml_results.append({
                            'comment_id': comment.comment_id,
                            'comment_text': comment.text[:50] + "..." if len(comment.text) > 50 else comment.text,
                            'status': final_status,
                            'score': score,
                            'confidence': ml_result['confidence']
                        })
                        
                        # Schedule deletion for compliant comments (5-minute delay)
                        if final_status == 'compliant' and score >= 0.6:
                            # Check if comment is already scheduled for deletion
                            if comment.scheduled_deletion_at is None:
                                # Schedule deletion in 5 minutes
                                from datetime import timedelta
                                scheduled_time = timezone.now() + timedelta(minutes=5)
                                comment.scheduled_deletion_at = scheduled_time
                                comment.save()
                                
                                print(f"SCHEDULED compliant comment {comment.comment_id} for deletion at {scheduled_time.strftime('%H:%M:%S')} (score: {ml_result['compliance_score']:.2f})")
                                compliant_comment_ids.append(comment.comment_id)
                            else:
                                print(f"ALREADY SCHEDULED comment {comment.comment_id} for deletion at {comment.scheduled_deletion_at.strftime('%H:%M:%S')}")
                        else:
                            # Clear any existing scheduled deletion if compliance changed
                            if comment.scheduled_deletion_at is not None:
                                comment.scheduled_deletion_at = None
                                comment.save()
                                print(f"CANCELLED scheduled deletion for comment {comment.comment_id} - no longer compliant")
                            print(f"KEEPING comment {comment.comment_id} - {ml_result['prediction']} (score: {ml_result['compliance_score']:.2f})")
                    
                    except Exception as e:
                        print(f"Error: ML compliance check failed for comment {comment.comment_id}: {e}")
                        print(f"Exception type: {type(e).__name__}")
                        print(f"Exception traceback: {e}")
                        # Keep comment with pending status on ML failure
                        comment.compliance_status = 'pending'
                        comment.save()
            
            print(f"DEBUG: Updating paragraph {paragraph_id} text in database...")
            # Update paragraph text in database
            paragraph.text = new_text
            # Clear html_content so the plain text will be displayed
            paragraph.html_content = ""
            paragraph.save()
            
            print(f"DEBUG: Updating paragraph {paragraph_id} text in DOCX file...")
            # Update paragraph text in DOCX file with file integrity check
            try:
                # Check if file exists and is a valid zip before attempting to modify it
                if not os.path.exists(document.file_path):
                    print(f"Warning: DOCX file not found at {document.file_path}, skipping DOCX update")
                else:
                    # Test if file is a valid zip/DOCX file
                    try:
                        with zipfile.ZipFile(document.file_path, 'r') as test_zip:
                            test_zip.testzip()  # This will raise an exception if the file is corrupted
                        print(f"DEBUG: DOCX file integrity check passed")
                        self.update_paragraph_in_docx(document.file_path, paragraph_id, new_text)
                        print(f"DEBUG: Successfully updated DOCX file")
                    except (zipfile.BadZipFile, zipfile.LargeZipFile) as zip_error:
                        print(f"Error: DOCX file is corrupted or not a valid zip file: {zip_error}")
                        print(f"Continuing with database update only...")
                        # Don't fail the entire request if DOCX update fails
            except Exception as file_error:
                print(f"Warning: Could not update DOCX file: {file_error}")
                print(f"Continuing with database update only...")
            
            print(f"DEBUG: Successfully completed EditParagraphView.put for paragraph {paragraph_id}")
            
            # Update response with ML compliance results and versioning info
            response_data = {
                'paragraph_id': paragraph.paragraph_id,
                'text': paragraph.text,
                'ml_compliance_results': ml_results,
                'document_id': document.id  # Include current document ID (may have changed due to auto-versioning)
            }
            
            # Add version information if auto-versioning occurred
            if version_created:
                response_data['version_created'] = True
                response_data['new_version_id'] = new_version_id
                response_data['new_version_number'] = new_version_number
                response_data['version_message'] = version_message
            elif version_message:  # Progress update without version creation
                response_data['version_message'] = version_message
            
            if compliant_comment_ids:
                response_data['scheduled_deletions'] = compliant_comment_ids
                if version_message:
                    response_data['message'] = f'{version_message}. {len(compliant_comment_ids)} compliant comment(s) scheduled for deletion in 5 minutes.'
                else:
                    response_data['message'] = f'Paragraph updated. {len(compliant_comment_ids)} compliant comment(s) scheduled for deletion in 5 minutes. {len(ml_results) - len(compliant_comment_ids)} comment(s) remain.'
            elif ml_results:
                if version_message:
                    response_data['message'] = f'{version_message}. {len(ml_results)} comment(s) checked - none were compliant enough for scheduled deletion.'
                else:
                    response_data['message'] = f'Paragraph updated. {len(ml_results)} comment(s) checked - none were compliant enough for scheduled deletion.'
            else:
                response_data['message'] = version_message or 'Paragraph updated.'
            
            return Response(response_data)
            
        except Document.DoesNotExist:
            print(f"ERROR: Document {document_id} not found")
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Paragraph.DoesNotExist:
            print(f"ERROR: Paragraph {paragraph_id} not found in document {document_id}")
            return Response({'error': 'Paragraph not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"CRITICAL ERROR in EditParagraphView.put: {e}")
            print(f"Exception type: {type(e).__name__}")
            import traceback
            traceback.print_exc()
            return Response({'error': f'Error updating paragraph: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def check_comment_compliance(self, original_text: str, comment_text: str, edited_text: str):
        """Check compliance using ML system (with fallback to basic system)"""
        try:
            # Try advanced ML system first
            if ML_FULL_SYSTEM_AVAILABLE and ML_DEPENDENCIES_AVAILABLE:
                ml_model = get_or_create_default_model()
                if ml_model is not None:
                    result = ml_model.predict(original_text, comment_text, edited_text)
                    return {
                        'prediction': result['prediction'],
                        'compliance_score': result['compliance_score'],
                        'confidence': result['confidence'],
                        'model_type': 'advanced_ml'
                    }
            
            # Fallback to basic system
            basic_model = get_basic_compliance_model()
            result = basic_model.predict(original_text, comment_text, edited_text)
            return {
                'prediction': result['prediction'],
                'compliance_score': result['compliance_score'],
                'confidence': result['confidence'],
                'model_type': 'basic'
            }
            
        except Exception as e:
            print(f"Warning: ML compliance check failed: {e}")
            # Return safe default on error
            return {
                'prediction': 'pending',
                'compliance_score': 0.0,
                'confidence': 0.0,
                'model_type': 'error_fallback'
            }
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Paragraph.DoesNotExist:
            return Response({'error': 'Paragraph not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Error editing paragraph: {e}")
            return Response({'error': f'Error editing paragraph: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def update_paragraph_in_docx(self, file_path, paragraph_id, new_text):
        try:
            # Debug: Check if method exists
            if not hasattr(self, '_recreate_docx_with_proper_xml_formatting'):
                print(f"ERROR: {self.__class__.__name__} does not have _recreate_docx_with_proper_xml_formatting method")
                print(f"MRO: {[cls.__name__ for cls in self.__class__.__mro__]}")
                # Fallback without XML formatting
                raise AttributeError("XML formatting method not available")
            
            # Create a backup
            backup_path = file_path + '.backup'
            shutil.copy2(file_path, backup_path)
            
            # Extract the DOCX
            temp_dir = file_path + '_temp'
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Update document.xml
            document_path = os.path.join(temp_dir, 'word', 'document.xml')
            
            if not os.path.exists(document_path):
                raise Exception("Document.xml not found")
            
            ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            
            tree = ET.parse(document_path)
            root = tree.getroot()
            
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Find and update the target paragraph
            paragraphs = root.findall('.//w:p', namespaces)
            paragraph_counter = 0
            
            for para in paragraphs:
                # Check if paragraph has text content
                para_text = ""
                for t in para.findall('.//w:t', namespaces):
                    if t.text:
                        para_text += t.text
                
                if para_text.strip():
                    paragraph_counter += 1
                    
                    if paragraph_counter == paragraph_id:
                        # Clear existing text elements but preserve structure
                        runs = para.findall('.//w:r', namespaces)
                        
                        # Remove all text elements from runs
                        for run in runs:
                            text_elements = run.findall('.//w:t', namespaces)
                            for text_elem in text_elements:
                                run.remove(text_elem)
                        
                        # If no runs exist, create one
                        if not runs:
                            new_run = ET.SubElement(para, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                            runs = [new_run]
                        
                        # Add new text to the first run
                        if new_text.strip():
                            first_run = runs[0]
                            new_text_elem = ET.SubElement(first_run, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                            new_text_elem.text = new_text
                            
                            if new_text != new_text.strip():
                                new_text_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        
                        break
            
            tree.write(document_path, encoding='utf-8', xml_declaration=True)
            
            # Recreate the DOCX file with proper XML formatting for ALL files
            self._recreate_docx_with_proper_xml_formatting(file_path, temp_dir)
            
            shutil.rmtree(temp_dir)
            os.remove(backup_path)
            
        except Exception as e:
            # Restore backup if something went wrong
            if os.path.exists(backup_path):
                shutil.copy2(backup_path, file_path)
                os.remove(backup_path)
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            raise e


class AddParagraphView(XMLFormattingMixin, APIView):
    def post(self, request):
        document_id = request.data.get('document_id')
        text = request.data.get('text', '')
        position = request.data.get('position')  # Optional: insert at specific position
        
        if not document_id:
            return Response({'error': 'Document ID is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Allow empty text for new paragraphs
        if text is None:
            text = ''
        
        try:
            document = Document.objects.get(id=document_id)
            
            # Determine the new paragraph ID
            if position and position > 0:
                new_paragraph_id = position
                # Update existing paragraphs with IDs >= position
                paragraphs_to_update = Paragraph.objects.filter(
                    document=document, 
                    paragraph_id__gte=position
                ).order_by('-paragraph_id')  # Order by descending to avoid conflicts
                
                for para in paragraphs_to_update:
                    para.paragraph_id += 1
                    para.save()
            else:
                # Add at the end
                last_paragraph = Paragraph.objects.filter(document=document).order_by('-paragraph_id').first()
                new_paragraph_id = (last_paragraph.paragraph_id + 1) if last_paragraph else 1
            
            # Create paragraph in database
            paragraph = Paragraph.objects.create(
                document=document,
                paragraph_id=new_paragraph_id,
                text=text
            )
            
            # Add paragraph to DOCX file
            self.add_paragraph_to_docx(document.file_path, new_paragraph_id, text, position)
            
            return Response({
                'paragraph_id': paragraph.paragraph_id,
                'text': paragraph.text,
                'message': 'Paragraph added successfully'
            }, status=status.HTTP_201_CREATED)
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Error adding paragraph: {e}")
            import traceback
            traceback.print_exc()  # for debigging
            return Response({'error': f'Error adding paragraph: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def add_paragraph_to_docx(self, file_path, paragraph_id, text, position=None):
        backup_path = None
        temp_dir = None
        
        try:
            # Create a backup
            backup_path = file_path + '.backup'
            shutil.copy2(file_path, backup_path)
            
            # Extract the DOCX
            temp_dir = file_path + '_temp'
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Update document.xml
            document_path = os.path.join(temp_dir, 'word', 'document.xml')
            
            if not os.path.exists(document_path):
                raise Exception(f"Document.xml not found at {document_path}")
            
            # Register namespace to preserve XML structure
            ET.register_namespace('', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            
            tree = ET.parse(document_path)
            root = tree.getroot()
            
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Create new paragraph element with proper namespace
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            new_para = ET.Element(f'{w_ns}p')
            new_run = ET.SubElement(new_para, f'{w_ns}r')
            new_text_elem = ET.SubElement(new_run, f'{w_ns}t')
            new_text_elem.text = text if text.strip() else ' '  # Ensure at least a space
            
            # Find the body element
            body = root.find('.//w:body', namespaces)
            if body is None:
                raise Exception("Document body not found")
            
            if position and position > 0:
                # Insert at specific position
                all_paragraphs = body.findall('w:p', namespaces)
                
                # Count only non-empty paragraphs to match our numbering system
                non_empty_count = 0
                insert_index = len(all_paragraphs)  # Default to end
                
                for i, para in enumerate(all_paragraphs):
                    # Check if paragraph has text content
                    para_text = ""
                    for t in para.findall('.//w:t', namespaces):
                        if t.text:
                            para_text += t.text
                    
                    if para_text.strip():
                        non_empty_count += 1
                        
                    if non_empty_count == position - 1:
                        insert_index = i + 1
                        break
                
                body.insert(insert_index, new_para)
            else:
                # Add at the end
                body.append(new_para)
            
            tree.write(document_path, encoding='utf-8', xml_declaration=True)
            
            # Recreate the DOCX file with proper XML formatting
            self._recreate_docx_with_proper_xml_formatting(file_path, temp_dir)
            
            # Cleanup
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            if backup_path and os.path.exists(backup_path):
                os.remove(backup_path)
            
        except Exception as e:
            # Restore backup if something went wrong
            if backup_path and os.path.exists(backup_path):
                shutil.copy2(backup_path, file_path)
                os.remove(backup_path)
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            raise e


@method_decorator(csrf_exempt, name='dispatch')
class DeleteParagraphView(XMLFormattingMixin, APIView):
    def delete(self, request):
        # Parse JSON data using DRF's request.data instead of raw request.body
        try:
            data = request.data
            document_id = data.get('document_id')
            paragraph_id = data.get('paragraph_id')
        except (AttributeError, ValueError) as e:
            return Response({'error': 'Invalid JSON in request data'}, status=status.HTTP_400_BAD_REQUEST)
        
        if not all([document_id, paragraph_id]):
            return Response({'error': 'Missing required fields'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            document = Document.objects.get(id=document_id)
            paragraph = Paragraph.objects.get(document=document, paragraph_id=paragraph_id)
            
            # Check if this is the last paragraph
            total_paragraphs = Paragraph.objects.filter(document=document).count()
            if total_paragraphs <= 1:
                return Response({'error': 'Cannot delete the last paragraph'}, status=status.HTTP_400_BAD_REQUEST)
            
            # Delete associated comments first
            deleted_comments = Comment.objects.filter(paragraph=paragraph)
            comment_count = deleted_comments.count()
            deleted_comments.delete()
            
            # Delete paragraph from DOCX file first (so we can restore if DB update fails)
            self.delete_paragraph_from_docx(document.file_path, paragraph_id)
            
            # Delete paragraph from database
            paragraph.delete()
            
            # Update paragraph IDs for paragraphs that come after the deleted one
            paragraphs_to_update = Paragraph.objects.filter(
                document=document, 
                paragraph_id__gt=paragraph_id
            ).order_by('paragraph_id')
            
            for para in paragraphs_to_update:
                para.paragraph_id -= 1
                para.save()
            
            # Also update comment references
            comments_to_update = Comment.objects.filter(
                document=document,
                paragraph__paragraph_id__gt=paragraph_id - 1  # After the shift
            )
            
            for comment in comments_to_update:
                # Find the paragraph with the updated ID
                try:
                    comment.paragraph = Paragraph.objects.get(
                        document=document,
                        paragraph_id=comment.paragraph.paragraph_id
                    )
                    comment.save()
                except Paragraph.DoesNotExist:
                    comment.delete()  # Remove orphaned comments
            
            return Response({
                'message': 'Paragraph deleted successfully',
                'deleted_comments': comment_count,
                'updated_paragraphs': paragraphs_to_update.count()
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Paragraph.DoesNotExist:
            return Response({'error': 'Paragraph not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Error deleting paragraph: {e}")
            import traceback
            traceback.print_exc()
            return Response({'error': f'Error deleting paragraph: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def delete_paragraph_from_docx(self, file_path, paragraph_id):
        backup_path = None
        temp_dir = None
        
        try:
            # Create a backup
            backup_path = file_path + '.backup'
            shutil.copy2(file_path, backup_path)
            
            # Extract the DOCX
            temp_dir = file_path + '_temp'
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
                
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Update document.xml
            document_path = os.path.join(temp_dir, 'word', 'document.xml')
            
            if not os.path.exists(document_path):
                raise Exception(f"Document.xml not found at {document_path}")
            
            # Register namespaces
            ET.register_namespace('', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            
            tree = ET.parse(document_path)
            root = tree.getroot()
            
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Find the body element first
            body = root.find('.//w:body', namespaces)
            if body is None:
                raise Exception("Document body not found")
            
            # Find and delete the target paragraph
            paragraphs = body.findall('w:p', namespaces)  # Direct children of body
            paragraph_counter = 0
            paragraph_deleted = False
            
            for i, para in enumerate(paragraphs):
                # Check if paragraph has text content
                para_text = ""
                for t in para.findall('.//w:t', namespaces):
                    if t.text:
                        para_text += t.text
                
                if para_text.strip():
                    paragraph_counter += 1
                    
                    if paragraph_counter == paragraph_id:
                        # Remove this paragraph from the body
                        body.remove(para)
                        paragraph_deleted = True
                        print(f"Deleted paragraph {paragraph_id} at position {i}")
                        break
            
            if not paragraph_deleted:
                raise Exception(f"Paragraph {paragraph_id} not found in document")
            
            # Write the updated XML
            tree.write(document_path, encoding='utf-8', xml_declaration=True)
            
            # Recreate the DOCX file with proper XML formatting
            self._recreate_docx_with_proper_xml_formatting(file_path, temp_dir)
            
            # Cleanup
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            if backup_path and os.path.exists(backup_path):
                os.remove(backup_path)
            
        except Exception as e:
            # Restore backup if something went wrong
            if backup_path and os.path.exists(backup_path):
                shutil.copy2(backup_path, file_path)
                os.remove(backup_path)
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            raise e
        
class AddCommentView(XMLFormattingMixin, APIView):
    def post(self, request):
        document_id = request.data.get('document_id')
        paragraph_id = request.data.get('paragraph_id')
        author = request.data.get('author', 'Anonymous')
        text = request.data.get('text', '')
        
        if not all([document_id, paragraph_id, text]):
            return Response({'error': 'Missing required fields'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            document = Document.objects.get(id=document_id)
            paragraph = Paragraph.objects.get(document=document, paragraph_id=paragraph_id)
            
            # Get next comment ID
            existing_comments = Comment.objects.filter(document=document)
            next_comment_id = max([c.comment_id for c in existing_comments] + [0]) + 1
            
            # Create comment in database
            comment = Comment.objects.create(
                document=document,
                paragraph=paragraph,
                comment_id=next_comment_id,
                author=author,
                text=text
            )
            
            # Update document status to 'commented' if it was 'original'
            document.update_status_based_on_comments()
            
            # Try to add comment to DOCX file
            docx_success = True
            docx_error = None
            try:
                self.add_comment_to_docx(document.file_path, paragraph_id, next_comment_id, author, text)
            except Exception as docx_e:
                docx_success = False
                docx_error = str(docx_e)
                print(f"Warning: Could not add comment to DOCX file: {docx_e}")
            
            return Response({
                'id': comment.comment_id,
                'comment_id': comment.comment_id,
                'author': comment.author,
                'text': comment.text,
                'paragraph_id': paragraph.paragraph_id,
                'docx_success': docx_success,
                'docx_error': docx_error if not docx_success else None
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Paragraph.DoesNotExist:
            return Response({'error': 'Paragraph not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Error adding comment: {e}")
            return Response({'error': f'Error adding comment: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def add_comment_to_docx(self, file_path, paragraph_id, comment_id, author, text):
        try:
            # Create a backup
            backup_path = file_path + '.backup'
            shutil.copy2(file_path, backup_path)
            
            # Extract the DOCX
            temp_dir = file_path + '_temp'
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            namespaces = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            }
            
            # Register namespace
            ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            
            # Update or create comments.xml
            comments_path = os.path.join(temp_dir, 'word', 'comments.xml')
            
            if os.path.exists(comments_path):
                # Load existing comments
                tree = ET.parse(comments_path)
                root = tree.getroot()
            else:
                # Create new comments.xml
                root = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comments')
                tree = ET.ElementTree(root)
                
                # Ensure word directory exists
                os.makedirs(os.path.join(temp_dir, 'word'), exist_ok=True)
            
            # Create new comment element
            comment_elem = ET.SubElement(root, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment')
            comment_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
            comment_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', author)
            comment_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', datetime.now().isoformat())
            
            # Add comment text
            p_elem = ET.SubElement(comment_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            r_elem = ET.SubElement(p_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
            t_elem = ET.SubElement(r_elem, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
            t_elem.text = text
            
            # Write XML with proper formatting
            self._write_xml_with_proper_formatting(tree, comments_path)
            
            # Update document.xml to add comment reference
            self.add_comment_reference_to_document(temp_dir, paragraph_id, comment_id)
            
            # Update relationships if needed
            self.ensure_comments_relationship(temp_dir)
            
            # Ensure comments content type is registered
            self.ensure_comments_content_type(temp_dir)
            
            # Recreate the DOCX file with proper XML formatting for ALL files
            self._recreate_docx_with_proper_xml_formatting(file_path, temp_dir)
            
            shutil.rmtree(temp_dir)
            os.remove(backup_path)
            
        except Exception as e:
            # Restore backup if something went wrong
            if os.path.exists(backup_path):
                shutil.copy2(backup_path, file_path)
                os.remove(backup_path)
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
            raise e

    def add_comment_reference_to_document(self, temp_dir, paragraph_id, comment_id):
        document_path = os.path.join(temp_dir, 'word', 'document.xml')
        
        if not os.path.exists(document_path):
            return
        
        ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
        
        tree = ET.parse(document_path)
        root = tree.getroot()
        
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Find the target paragraph
        paragraphs = root.findall('.//w:p', namespaces)
        paragraph_counter = 0
        
        for para in paragraphs:
            # Check if paragraph has text content
            para_text = ""
            for t in para.findall('.//w:t', namespaces):
                if t.text:
                    para_text += t.text
            
            if para_text.strip():
                paragraph_counter += 1
                
                if paragraph_counter == paragraph_id:
                    # Find the first run in the paragraph
                    first_run = para.find('.//w:r', namespaces)
                    if first_run is not None:
                        # Add comment range start
                        comment_start = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeStart')
                        comment_start.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                        para.insert(0, comment_start)
                        
                        # Add comment range end
                        comment_end = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentRangeEnd')
                        comment_end.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                        para.append(comment_end)
                        
                        # Add comment reference
                        comment_ref_run = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')
                        comment_ref = ET.SubElement(comment_ref_run, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}commentReference')
                        comment_ref.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id', str(comment_id))
                        para.append(comment_ref_run)
                    
                    break
        
        tree.write(document_path, encoding='utf-8', xml_declaration=True)

    def ensure_comments_relationship(self, temp_dir):
        rels_path = os.path.join(temp_dir, 'word', '_rels', 'document.xml.rels')
        
        if not os.path.exists(rels_path):
            # Create the _rels directory and file if it doesn't exist
            os.makedirs(os.path.dirname(rels_path), exist_ok=True)
            
            # Create basic relationships file
            rels_root = ET.Element('Relationships')
            rels_root.set('xmlns', 'http://schemas.openxmlformats.org/package/2006/relationships')
        else:
            tree = ET.parse(rels_path)
            rels_root = tree.getroot()
        
        comments_rel_exists = False
        for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
            if rel.get('Target') == 'comments.xml':
                comments_rel_exists = True
                break
        
        if not comments_rel_exists:
            rel_elem = ET.SubElement(rels_root, '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            # Generate a unique relationship ID
            existing_ids = [rel.get('Id', '') for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')]
            rel_id = f"rId{max([int(rid[3:]) for rid in existing_ids if rid.startswith('rId') and rid[3:].isdigit()] + [0]) + 1}"
            
            rel_elem.set('Id', rel_id)
            rel_elem.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')
            rel_elem.set('Target', 'comments.xml')
            
            # Save the relationships file
            ET.register_namespace('', 'http://schemas.openxmlformats.org/package/2006/relationships')
            tree = ET.ElementTree(rels_root)
            self._write_xml_with_proper_formatting(tree, rels_path)

    def ensure_comments_content_type(self, temp_dir):
        """Ensure comments.xml is registered in [Content_Types].xml"""
        content_types_path = os.path.join(temp_dir, '[Content_Types].xml')
        
        if not os.path.exists(content_types_path):
            return
        
        try:
            tree = ET.parse(content_types_path)
            root = tree.getroot()
            
            # Check if comments content type already exists
            namespaces = {'ct': 'http://schemas.openxmlformats.org/package/2006/content-types'}
            existing = root.find(".//ct:Override[@PartName='/word/comments.xml']", namespaces)
            
            if existing is None:
                # Add the comments content type
                override_elem = ET.SubElement(root, '{http://schemas.openxmlformats.org/package/2006/content-types}Override')
                override_elem.set('PartName', '/word/comments.xml')
                override_elem.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
                
                # Save the content types file
                ET.register_namespace('', 'http://schemas.openxmlformats.org/package/2006/content-types')
                self._write_xml_with_proper_formatting(tree, content_types_path)
                
        except Exception as e:
            print(f"Error updating content types: {e}")


@method_decorator(csrf_exempt, name='dispatch')
class DeleteCommentView(XMLFormattingMixin, APIView):
    def delete(self, request):
        # Parse JSON data from request body for DELETE requests
        try:
            import json
            data = json.loads(request.body.decode('utf-8'))
            document_id = data.get('document_id')
            comment_id = data.get('comment_id')
        except (json.JSONDecodeError, UnicodeDecodeError, AttributeError):
            # Fallback to request.data if JSON parsing fails
            try:
                data = request.data
                document_id = data.get('document_id')
                comment_id = data.get('comment_id')
            except (AttributeError, ValueError) as e:
                return Response({'error': 'Invalid JSON in request data'}, status=status.HTTP_400_BAD_REQUEST)
        
        if document_id is None or comment_id is None:
            error_msg = f'Missing required fields: document_id={document_id}, comment_id={comment_id}'
            return Response({'error': error_msg}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            document = Document.objects.get(id=document_id)
            comment = Comment.objects.get(document=document, comment_id=comment_id)
            
            # Delete comment from DOCX file first
            self.delete_comment_from_docx(document.file_path, comment_id)
            
            # Delete comment from database
            comment.delete()
            
            return Response({
                'message': 'Comment deleted successfully',
                'comment_id': comment_id
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Comment.DoesNotExist:
            return Response({'error': 'Comment not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Error deleting comment: {e}")
            return Response({'error': f'Error deleting comment: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class ListDocumentsView(APIView):
    def get(self, request):
        # Show all documents in both interfaces
        documents = Document.objects.all()
            
        # Add comment count and order by upload date
        documents = documents.annotate(comment_count=Count('comments')).order_by('-uploaded_at')
        
        documents_data = []
        for doc in documents:
            documents_data.append({
                'id': doc.id,
                'filename': doc.filename,
                'uploaded_at': doc.uploaded_at.isoformat(),
                'comment_count': doc.comment_count,
                # Version information
                'version_number': doc.version_number,
                'version_status': doc.version_status,
                'is_original': doc.version_number == 1,
                'parent_document_id': doc.parent_document.id if doc.parent_document else None,
                'base_document_id': doc.base_document.id if doc.base_document else doc.id,
                'created_from_comments': doc.created_from_comments,
                'version_notes': doc.version_notes
            })
        
        return Response(documents_data)

class ExportDocumentView(APIView):
    def get(self, request, document_id):
        try:
            document = Document.objects.get(id=document_id)
            print(f"Exporting document {document_id}: {document.filename}")
            print(f"File path: {document.file_path}")
            print(f"Version number: {document.version_number}")
            print(f"Version status: {document.version_status}")
            print(f"Is base document: {document.base_document is None}")
            
            # Check if file exists and get basic info
            if os.path.exists(document.file_path):
                file_size = os.path.getsize(document.file_path)
                file_mtime = os.path.getmtime(document.file_path)
                print(f"File size: {file_size} bytes")
                print(f"Last modified: {datetime.fromtimestamp(file_mtime)}")
                
                # Quick check if file is valid DOCX
                try:
                    with zipfile.ZipFile(document.file_path, 'r') as test_zip:
                        file_list = test_zip.namelist()
                        print(f"DOCX contains {len(file_list)} files")
                        if 'word/document.xml' in file_list:
                            print("✓ Valid DOCX structure detected")
                        else:
                            print("⚠ Missing document.xml - file may be corrupted")
                except Exception as zip_error:
                    print(f"⚠ DOCX file validation failed: {zip_error}")
            else:
                print(f"❌ File not found at: {document.file_path}")
                
            # Check if we should rebuild the DOCX from database (sync option)
            sync_with_db = request.GET.get('sync', '').lower() in ['true', '1', 'yes']
            if sync_with_db:
                print("🔄 Rebuilding DOCX from database content...")
                try:
                    self._rebuild_docx_from_database(document)
                    print("✅ DOCX rebuilt successfully from database")
                except Exception as rebuild_error:
                    print(f"❌ Failed to rebuild DOCX: {rebuild_error}")
                    # Continue with export even if rebuild fails
                
            # Generate version-aware filename
            base_filename = document.filename
            if not base_filename.lower().endswith('.docx'):
                base_filename += '.docx'
            
            # Remove .docx extension to insert version info
            name_without_ext = base_filename[:-5] if base_filename.lower().endswith('.docx') else base_filename
            
            # Create export filename based on version
            if document.version_number > 1:
                export_filename = f"{name_without_ext}_v{document.version_number}.docx"
            else:
                # For version 1, we can still indicate it's been processed/updated
                if document.version_status in ['edited', 'commented']:
                    export_filename = f"{name_without_ext}_updated.docx"
                else:
                    export_filename = base_filename
            
            print(f"Export filename: {export_filename}")
            
            if not document.file_path:
                return Response({'error': 'No file path saved for document'}, status=status.HTTP_404_NOT_FOUND)
                return Response({'error': 'No file path saved for document'}, status=status.HTTP_404_NOT_FOUND)
            
            if os.path.exists(document.file_path):
                try:
                    file = open(document.file_path, 'rb')
                    response = FileResponse(
                        file,
                        as_attachment=True,
                        filename=export_filename,
                        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    return response
                except PermissionError as e:
                    print(f"Permission error: {str(e)}")
                    return Response({'error': f'Permission denied: {str(e)}'}, status=status.HTTP_403_FORBIDDEN)
                except Exception as e:
                    print(f"File open error: {str(e)}")
                    return Response({'error': f'Error opening file: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            else:
                print(f"File not found at path: {document.file_path}")
                return Response({
                    'error': f'File not found at path: {document.file_path}'
                }, status=status.HTTP_404_NOT_FOUND)
                
        except Document.DoesNotExist:
            print(f"Document {document_id} not found in database")
            return Response({'error': 'Document not found in database'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Unexpected error in export: {str(e)}")
            return Response({'error': f'Export error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def _rebuild_docx_from_database(self, document):
        """Rebuild the DOCX file content from the current database state"""
        from docx import Document as DocxDocument
        
        # Create a new DOCX document
        new_doc = DocxDocument()
        
        # Clear default paragraph
        if new_doc.paragraphs:
            p = new_doc.paragraphs[0]
            p.clear()
        
        # Add paragraphs from database
        paragraphs = document.paragraphs.all().order_by('paragraph_id')
        for para in paragraphs:
            if para.html_content and para.html_content.strip():
                # Handle HTML content - extract text for now (could be enhanced for formatting)
                from html import unescape
                import re
                text_content = re.sub('<[^<]+?>', '', para.html_content)
                text_content = unescape(text_content).strip()
            else:
                text_content = para.text or ''
            
            if text_content:
                new_doc.add_paragraph(text_content)
            else:
                new_doc.add_paragraph('')  # Keep empty paragraphs
        
        # Save to the document's file path
        backup_path = document.file_path + '.backup'
        if os.path.exists(document.file_path):
            shutil.copy2(document.file_path, backup_path)
        
        try:
            new_doc.save(document.file_path)
            # Remove backup if successful
            if os.path.exists(backup_path):
                os.remove(backup_path)
        except Exception as e:
            # Restore backup if save failed
            if os.path.exists(backup_path):
                shutil.copy2(backup_path, document.file_path)
                os.remove(backup_path)
            raise e


class GetDocumentView(APIView):
    def get(self, request, document_id):
        try:
            document = Document.objects.get(id=document_id)
            
            paragraphs_data = []
            for para in document.paragraphs.all().order_by('paragraph_id'):
                para_data = {
                    'id': para.paragraph_id,
                    'text': para.text,
                    'html_content': para.html_content,
                    'has_images': para.has_images
                }
                
                # Add image information if paragraph has images
                if para.has_images:
                    images = []
                    for para_img in para.paragraph_images.all():
                        images.append({
                            'id': para_img.document_image.id,
                            'filename': para_img.document_image.filename,
                            'image_id': para_img.document_image.image_id,
                            'position': para_img.position_in_paragraph
                        })
                    para_data['images'] = images
                
                paragraphs_data.append(para_data)
            
            comments_data = []
            for comment in document.comments.all():
                comments_data.append({
                    'id': comment.comment_id,
                    'author': comment.author,
                    'text': comment.text,
                    'paragraph_id': comment.paragraph.paragraph_id
                })
            
            return Response({
                'document_id': document.id,
                'filename': document.filename,
                'version_number': document.version_number,
                'version_status': document.version_status,
                'created_from_comments': document.created_from_comments,
                'parent_document_id': document.parent_document.id if document.parent_document else None,
                'base_document_id': document.base_document.id if document.base_document else document.id,
                'version_notes': document.version_notes,
                'edited_commented_paragraphs': document.edited_commented_paragraphs,
                'uploaded_at': document.uploaded_at.isoformat(),
                'paragraphs': paragraphs_data,
                'comments': comments_data
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)


class ServeImageView(APIView):
    def get(self, request, image_id):
        """Serve document images"""
        try:
            image = DocumentImage.objects.get(id=image_id)
            
            if os.path.exists(image.file_path):
                response = FileResponse(
                    open(image.file_path, 'rb'),
                    content_type=image.content_type
                )
                return response
            else:
                return Response({'error': 'Image file not found'}, status=status.HTTP_404_NOT_FOUND)
                
        except DocumentImage.DoesNotExist:
            return Response({'error': 'Image not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({'error': f'Error serving image: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# ============================================================================
# ML COMPLIANCE CHECKING VIEWS
# ============================================================================

import time
from django.utils import timezone
from .basic_ml_compliance import get_basic_compliance_model

# Try to import full ML system
try:
    from .ml_compliance import get_or_create_default_model, ML_DEPENDENCIES_AVAILABLE
    ML_FULL_SYSTEM_AVAILABLE = True
except ImportError:
    ML_FULL_SYSTEM_AVAILABLE = False
    ML_DEPENDENCIES_AVAILABLE = False


class CheckEditComplianceRealTimeView(APIView):
    """
    Real-time ML compliance checking for live editing feedback
    This endpoint is called while the user is typing to provide instant feedback
    """
    
    def post(self, request):
        # Extract input data
        paragraph_id = request.data.get('paragraph_id')
        document_id = request.data.get('document_id')
        current_text = request.data.get('current_text', '')
        
        if not all([paragraph_id, document_id, current_text]):
            return Response({
                'error': 'Missing required fields: paragraph_id, document_id, current_text'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Get document and paragraph
            document = Document.objects.get(id=document_id)
            paragraph = Paragraph.objects.get(document=document, paragraph_id=paragraph_id)
            
            # Get original text and all comments for this paragraph
            original_text = paragraph.text or ""
            comments = Comment.objects.filter(paragraph=paragraph)
            
            if not comments.exists():
                return Response({
                    'message': 'No comments found for this paragraph',
                    'compliance_results': []
                })
            
            # Check compliance against all comments in real-time
            compliance_results = []
            overall_status = 'compliant'  # Start optimistic
            
            for comment in comments:
                try:
                    # Use the same ML checking method as edit paragraph
                    ml_result = self.check_comment_compliance_realtime(original_text, comment.text, current_text)
                    
                    result_data = {
                        'comment_id': comment.comment_id,
                        'comment_text': comment.text,
                        'status': ml_result['prediction'],
                        'score': ml_result['compliance_score'],
                        'confidence': ml_result['confidence'],
                        'model_type': ml_result['model_type'],
                        'scheduled_deletion_at': comment.scheduled_deletion_at.isoformat() if comment.scheduled_deletion_at else None,
                        'is_scheduled_for_deletion': comment.scheduled_deletion_at is not None
                    }
                    
                    compliance_results.append(result_data)
                    
                    # Determine overall status (most restrictive)
                    if ml_result['prediction'] == 'non_compliant':
                        overall_status = 'non_compliant'
                    elif ml_result['prediction'] == 'partial' and overall_status == 'compliant':
                        overall_status = 'partial'
                
                except Exception as e:
                    print(f"Real-time compliance check failed for comment {comment.comment_id}: {e}")
                    compliance_results.append({
                        'comment_id': comment.comment_id,
                        'comment_text': comment.text,
                        'status': 'error',
                        'score': 0.0,
                        'confidence': 0.0,
                        'model_type': 'error',
                        'scheduled_deletion_at': comment.scheduled_deletion_at.isoformat() if comment.scheduled_deletion_at else None,
                        'is_scheduled_for_deletion': comment.scheduled_deletion_at is not None
                    })
            
            # Calculate overall compliance
            if compliance_results:
                avg_score = sum(r['score'] for r in compliance_results if r['score'] > 0) / len([r for r in compliance_results if r['score'] > 0])
            else:
                avg_score = 0.0
            
            return Response({
                'overall_status': overall_status,
                'overall_score': avg_score,
                'compliance_results': compliance_results,
                'total_comments': len(compliance_results),
                'can_auto_delete': overall_status == 'compliant' and avg_score >= 0.6
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Paragraph.DoesNotExist:
            return Response({'error': 'Paragraph not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'error': f'Error during real-time compliance check: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    def check_comment_compliance_realtime(self, original_text: str, comment_text: str, edited_text: str):
        """Check compliance using ML system - optimized for real-time use with consistent thresholds"""
        try:
            # Try advanced ML system first
            if ML_FULL_SYSTEM_AVAILABLE and ML_DEPENDENCIES_AVAILABLE:
                ml_model = get_or_create_default_model()
                if ml_model is not None:
                    result = ml_model.predict(original_text, comment_text, edited_text)
                    score = result['compliance_score']
                    
                    # Use consistent classification thresholds
                    if score >= 0.6:
                        final_status = 'compliant'
                    elif score >= 0.3:
                        final_status = 'partial'
                    else:
                        final_status = 'non_compliant'
                    
                    return {
                        'prediction': final_status,
                        'compliance_score': score,
                        'confidence': result['confidence'],
                        'model_type': 'advanced_ml'
                    }
            
            # Fallback to basic system
            basic_model = get_basic_compliance_model()
            result = basic_model.predict(original_text, comment_text, edited_text)
            score = result['compliance_score']
            
            # Use consistent classification thresholds for basic model too
            if score >= 0.6:
                final_status = 'compliant'
            elif score >= 0.3:
                final_status = 'partial'
            else:
                final_status = 'non_compliant'
            
            return {
                'prediction': final_status,
                'compliance_score': score,
                'confidence': result['confidence'],
                'model_type': 'basic'
            }
            
        except Exception as e:
            print(f"Real-time ML compliance check failed: {e}")
            return {
                'prediction': 'pending',
                'compliance_score': 0.0,
                'confidence': 0.0,
                'model_type': 'error_fallback'
            }


class CheckEditComplianceView(APIView):
    """
    API endpoint to check if an edit complies with a comment
    """
    
    def post(self, request):
        # Extract input data
        original_text = request.data.get('original_text', '')
        comment_text = request.data.get('comment_text', '')
        edited_text = request.data.get('edited_text', '')
        
        if not all([original_text, comment_text, edited_text]):
            return Response({
                'error': 'Missing required fields: original_text, comment_text, edited_text'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Try to use full ML model first, fallback to basic if needed
            model_type = 'basic'
            
            if ML_FULL_SYSTEM_AVAILABLE and ML_DEPENDENCIES_AVAILABLE:
                try:
                    # Try to get the advanced ML model
                    ml_model = get_or_create_default_model()
                    if ml_model is not None:
                        # Record start time for performance tracking
                        start_time = time.time()
                        
                        # Make prediction using advanced ML model
                        prediction_result = ml_model.predict(original_text, comment_text, edited_text)
                        explanation = ml_model.explain_prediction(original_text, comment_text, edited_text)
                        
                        # Calculate processing time
                        processing_time = int((time.time() - start_time) * 1000)  # milliseconds
                        
                        response_data = {
                            'compliance_score': prediction_result['compliance_score'],
                            'prediction': prediction_result['prediction'],
                            'confidence': prediction_result['confidence'],
                            'explanation': {
                                'interpretation': explanation['interpretation'],
                                'top_features': explanation.get('top_features', [])
                            },
                            'processing_time_ms': processing_time,
                            'model_type': 'advanced_ml'
                        }
                        
                        return Response(response_data)
                except Exception as e:
                    print(f"Warning: Full ML system failed, falling back to basic: {e}")
            
            # Fallback to basic compliance model
            model = get_basic_compliance_model()
            
            # Record start time for performance tracking
            start_time = time.time()
            
            # Make prediction
            prediction_result = model.predict(original_text, comment_text, edited_text)
            
            # Get explanation
            explanation = model.explain_prediction(original_text, comment_text, edited_text)
            
            # Calculate processing time
            processing_time = int((time.time() - start_time) * 1000)  # milliseconds
            
            response_data = {
                'compliance_score': prediction_result['compliance_score'],
                'prediction': prediction_result['prediction'],
                'confidence': prediction_result['confidence'],
                'explanation': {
                    'interpretation': explanation['interpretation']
                },
                'processing_time_ms': processing_time,
                'model_type': 'basic'
            }
            
            return Response(response_data)
            
        except Exception as e:
            return Response({
                'error': f'Error during compliance check: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class CheckParagraphComplianceView(APIView):
    """
    Check compliance for a specific paragraph that was edited
    """
    
    def post(self, request):
        document_id = request.data.get('document_id')
        paragraph_id = request.data.get('paragraph_id')
        original_text = request.data.get('original_text', '')
        edited_text = request.data.get('edited_text', '')
        
        if not all([document_id, paragraph_id]):
            return Response({
                'error': 'Missing required fields: document_id, paragraph_id'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            document = Document.objects.get(id=document_id)
            paragraph = Paragraph.objects.get(document=document, paragraph_id=paragraph_id)
            
            # Get all comments for this paragraph
            comments = Comment.objects.filter(paragraph=paragraph)
            
            if not comments.exists():
                return Response({
                    'message': 'No comments found for this paragraph',
                    'compliance_results': []
                })
            
            # Use paragraph text if not provided
            if not original_text:
                original_text = paragraph.text
            if not edited_text:
                edited_text = paragraph.text
            
            # Check compliance against all comments
            compliance_results = []
            model = get_basic_compliance_model()
            
            for comment in comments:
                result = model.predict(original_text, comment.text, edited_text)
                explanation = model.explain_prediction(original_text, comment.text, edited_text)
                
                compliance_results.append({
                    'comment_id': comment.comment_id,
                    'comment_author': comment.author,
                    'comment_text': comment.text,
                    'compliance_score': result['compliance_score'],
                    'prediction': result['prediction'],
                    'confidence': result['confidence'],
                    'explanation': explanation['interpretation']
                })
            
            # Calculate overall compliance
            if compliance_results:
                avg_compliance = sum(r['compliance_score'] for r in compliance_results) / len(compliance_results)
                overall_prediction = 'compliant' if avg_compliance > 0.6 else 'partial' if avg_compliance > 0.3 else 'non_compliant'
            else:
                avg_compliance = 0.0
                overall_prediction = 'no_comments'
            
            return Response({
                'paragraph_id': paragraph_id,
                'overall_compliance_score': avg_compliance,
                'overall_prediction': overall_prediction,
                'compliance_results': compliance_results,
                'total_comments_checked': len(compliance_results)
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Paragraph.DoesNotExist:
            return Response({'error': 'Paragraph not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'error': f'Error checking paragraph compliance: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class MLModelStatusView(APIView):
    """Get information about the current ML model"""
    
    def get(self, request):
        try:
            print(f"DEBUG: ML_FULL_SYSTEM_AVAILABLE = {ML_FULL_SYSTEM_AVAILABLE}")
            print(f"DEBUG: ML_DEPENDENCIES_AVAILABLE = {ML_DEPENDENCIES_AVAILABLE}")
            
            # Check for advanced ML system first
            if ML_FULL_SYSTEM_AVAILABLE and ML_DEPENDENCIES_AVAILABLE:
                try:
                    print("DEBUG: Attempting to get advanced ML model...")
                    ml_model = get_or_create_default_model()
                    print(f"DEBUG: ML model result: {ml_model is not None}")
                    if ml_model is not None:
                        return Response({
                            'model_loaded': True,
                            'model_type': 'Advanced ML Classifier (RandomForest)',
                            'status': 'ready',
                            'ml_available': True,
                            'description': 'Full ML system with RandomForest classifier and 20+ features for accurate compliance prediction'
                        })
                except Exception as e:
                    print(f"DEBUG: Advanced ML system failed: {e}")
            
            # Fallback to basic system
            print("DEBUG: Using fallback basic system")
            return Response({
                'model_loaded': True,
                'model_type': 'Rule-based Basic Checker',
                'status': 'ready',
                'ml_available': ML_DEPENDENCIES_AVAILABLE,
                'description': 'Basic rule-based compliance checker (full ML dependencies available for upgrade)'
            })
            
        except Exception as e:
            print(f"ERROR: MLModelStatusView failed: {e}")
            return Response({
                'model_loaded': False,
                'model_type': 'Error',
                'status': 'error',
                'error': str(e)
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            

class CreateNewVersionView(XMLFormattingMixin, APIView):
    """
    Create a new version of a document based on processed comments.
    This is the core workflow: v1 + comments → v2
    """
    
    def post(self, request):
        """
        Create a new version by processing comments from the current version
        Expected data:
        {
            "document_id": int,  # Current version ID
            "version_notes": str,  # Optional notes about the changes
            "selected_comment_ids": [int],  # Optional: specific comments to process
        }
        """
        document_id = request.data.get('document_id')
        version_notes = request.data.get('version_notes', '')
        selected_comment_ids = request.data.get('selected_comment_ids', [])
        
        if not document_id:
            return Response({'error': 'Missing required field: document_id'}, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            # Get current version
            current_version = Document.objects.get(id=document_id)
            
            # Validate that current version has comments
            if not current_version.has_comments():
                return Response({
                    'error': 'Document has no comments to process'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Get base document (or use current as base if it's the original)
            base_doc = current_version.base_document or current_version
            
            # Calculate next version number
            next_version_number = base_doc.get_next_version_number()
            
            # Create new filename for the new version
            original_name = base_doc.filename
            name_parts = os.path.splitext(original_name)
            new_filename = f"{name_parts[0]}_v{next_version_number}{name_parts[1]}"
            
            # Copy the DOCX file to a new location
            original_path = current_version.file_path
            media_dir = os.path.dirname(original_path)
            new_file_path = os.path.join(media_dir, new_filename)
            
            import shutil
            shutil.copy2(original_path, new_file_path)
            
            # Create new document version
            new_version = Document.objects.create(
                filename=new_filename,
                file_path=new_file_path,
                is_editable=True,
                version_number=next_version_number,
                version_status='edited',
                base_document=base_doc,
                parent_document=current_version,
                created_from_comments=True,
                version_notes=version_notes
            )
            
            # Copy paragraphs from current version
            current_paragraphs = current_version.paragraphs.all().order_by('paragraph_id')
            paragraph_mapping = {}  # Map old paragraph IDs to new objects
            
            for old_paragraph in current_paragraphs:
                new_paragraph = Paragraph.objects.create(
                    document=new_version,
                    paragraph_id=old_paragraph.paragraph_id,
                    text=old_paragraph.text,
                    html_content=old_paragraph.html_content,
                    has_images=old_paragraph.has_images
                )
                paragraph_mapping[old_paragraph.paragraph_id] = new_paragraph
                
                # Copy paragraph images if any
                for para_image in old_paragraph.paragraph_images.all():
                    # Note: We're reusing the same DocumentImage objects
                    # since they're shared across versions
                    ParagraphImage.objects.create(
                        paragraph=new_paragraph,
                        document_image=para_image.document_image,
                        position_in_paragraph=para_image.position_in_paragraph
                    )
            
            # Get comments to process
            if selected_comment_ids:
                comments_to_process = current_version.comments.filter(
                    comment_id__in=selected_comment_ids
                )
            else:
                comments_to_process = current_version.comments.all()
            
            # Record which comments were processed
            processed_ids = list(comments_to_process.values_list('comment_id', flat=True))
            new_version.processed_comment_ids = processed_ids
            new_version.save()
            
            # Update current version status
            current_version.version_status = 'archived'
            current_version.save()
            
            # Parse the new document to update content
            from .docx_parser import EnhancedDocxParser
            parser = EnhancedDocxParser(new_file_path, new_version)
            paragraphs_data = parser.parse_document()
            
            return Response({
                'status': 'success',
                'message': f'Successfully created version {next_version_number}',
                'data': {
                    'new_version_id': new_version.id,
                    'new_version_number': next_version_number,
                    'filename': new_filename,
                    'processed_comments': len(processed_ids),
                    'comment_ids_processed': processed_ids,
                    'paragraphs_count': len(paragraphs_data),
                    'version_notes': version_notes
                }
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            print(f"Error creating new version: {e}")
            return Response({
                'error': f'Error creating new version: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class GetDocumentVersionsView(APIView):
    """
    Get all versions of a document chain
    """
    
    def get(self, request, document_id):
        try:
            document = Document.objects.get(id=document_id)
            base_doc = document.base_document or document
            
            versions = Document.objects.filter(
                Q(id=base_doc.id) | Q(base_document=base_doc)
            ).order_by('version_number')
            
            versions_data = []
            for version in versions:
                versions_data.append({
                    'id': version.id,
                    'version_number': version.version_number,
                    'filename': version.filename,
                    'version_status': version.version_status,
                    'created_from_comments': version.created_from_comments,
                    'comment_count': version.comments.count(),
                    'uploaded_at': version.uploaded_at.isoformat(),
                    'version_notes': version.version_notes,
                    'processed_comment_ids': version.processed_comment_ids,
                    'is_current': version.id == document.id
                })
            
            return Response({
                'base_document_id': base_doc.id,
                'current_version_id': document.id,
                'total_versions': len(versions_data),
                'versions': versions_data
            })
            
        except Document.DoesNotExist:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            return Response({
                'error': f'Error retrieving versions: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class DocumentVersionStatsView(APIView):
    """
    Get statistics about document versions and workflow
    """
    
    def get(self, request):
        try:
            total_documents = Document.objects.count()
            original_documents = Document.objects.filter(version_number=1).count()
            edited_versions = Document.objects.filter(version_number__gt=1).count()
            
            # Status distribution
            status_counts = {}
            for status_choice in Document.VERSION_STATUS_CHOICES:
                status_key = status_choice[0]
                status_counts[status_key] = Document.objects.filter(version_status=status_key).count()
            
            # Comments to processing stats
            commented_docs = Document.objects.filter(version_status='commented').count()
            docs_with_comments = Document.objects.filter(comments__isnull=False).distinct().count()
            
            return Response({
                'total_documents': total_documents,
                'original_documents': original_documents,
                'edited_versions': edited_versions,
                'status_distribution': status_counts,
                'workflow_stats': {
                    'commented_documents': commented_docs,
                    'documents_with_comments': docs_with_comments,
                    'ready_for_editing': commented_docs
                }
            })
            
        except Exception as e:
            return Response({
                'error': f'Error retrieving stats: {str(e)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
        except Exception as e:
            return Response({
                'error': f'Error getting model status: {str(e)}',
                'model_loaded': False,
                'status': 'error'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)